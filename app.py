import streamlit as st
from google import genai
from google.genai import types
from docx import Document
import random
import os
import io
import json
import warnings

# Menyembunyikan peringatan tidak penting dari library pendukung
warnings.filterwarnings("ignore")

# ==============================================================================
# 1. DATABASE HYBRID ADAPTER (AUTO-LOCAL JSON & AUTO-CLOUD STREAMLIT KV)
# ==============================================================================
EMAIL_ADMIN = "handoyoyy1@gmail.com"
NAMA_FILE_KEY = "api_keys.txt"
NAMA_FILE_DB = "database_users.json"  # Berkas cadangan otomatis khusus di localhost
NAMA_FILE_SESSION = "session_login.json"

def muat_database_kv():
    """Membaca data user dengan deteksi otomatis lingkungan kerja (Local/Cloud)."""
    # ------ JALUR 1: JIKA RUNNING DI STREAMLIT CLOUD (MENGGUNAKAN ST.KV) ------
    if hasattr(st, "kv"):
        if "db_users_master" not in st.kv:
            data_awal = {
                EMAIL_ADMIN: {
                    "nama": "Miftada Handoyo (Admin)",
                    "status": "Aktif", 
                    "nominal_transfer": 0, 
                    "kode_aktivasi": "ADMIN_ACCESS"
                }
            }
            st.kv["db_users_master"] = json.dumps(data_awal)
            return data_awal
        try:
            return json.loads(st.kv["db_users_master"])
        except:
            return {EMAIL_ADMIN: {"nama": "Miftada Handoyo (Admin)", "status": "Aktif", "nominal_transfer": 0, "kode_aktivasi": "ADMIN_ACCESS"}}
            
    # ------ JALUR 2: JIKA RUNNING DI LOCALHOST (FALLBACK KE FILE BERKAS JSON) ------
    else:
        if not os.path.exists(NAMA_FILE_DB):
            data_awal = {
                EMAIL_ADMIN: {
                    "nama": "Miftada Handoyo (Admin)",
                    "status": "Aktif", 
                    "nominal_transfer": 0, 
                    "kode_aktivasi": "ADMIN_ACCESS"
                }
            }
            with open(NAMA_FILE_DB, "w") as f:
                json.dump(data_awal, f, indent=4)
            return data_awal
        try:
            with open(NAMA_FILE_DB, "r") as f:
                return json.load(f)
        except:
            return {EMAIL_ADMIN: {"nama": "Miftada Handoyo (Admin)", "status": "Aktif", "nominal_transfer": 0, "kode_aktivasi": "ADMIN_ACCESS"}}

def simpan_database_kv(data_db):
    """Menyimpan pembaruan data user secara adaptif berdasarkan lokasi server."""
    # Jika di Cloud, simpan langsung ke memori st.kv
    if hasattr(st, "kv"):
        st.kv["db_users_master"] = json.dumps(data_db)
    # Jika di Local, tulis ke dalam file JSON fisik di hardisk
    else:
        with open(NAMA_FILE_DB, "w") as f:
            json.dump(data_db, f, indent=4)

# --- MANAGEMENT SESSION LOGIN ---
def simpan_session_login(email, nama):
    session_data = {"email": email, "nama": nama, "is_logged_in": True}
    with open(NAMA_FILE_SESSION, "w") as f:
        json.dump(session_data, f, indent=4)

def hapus_session_login():
    if os.path.exists(NAMA_FILE_SESSION):
        os.remove(NAMA_FILE_SESSION)

def cek_auto_login():
    if os.path.exists(NAMA_FILE_SESSION):
        try:
            with open(NAMA_FILE_SESSION, "r") as f:
                data = json.load(f)
                return data.get("email", ""), data.get("nama", ""), data.get("is_logged_in", False)
        except:
            return "", "", False
    return "", "", False

# Memuat status database dan session state saat aplikasi dijalankan
st.session_state.db_users = muat_database_kv()
saved_email, saved_name, is_saved_logged_in = cek_auto_login()

if "user_email" not in st.session_state: st.session_state.user_email = saved_email
if "user_name" not in st.session_state: st.session_state.user_name = saved_name
if "is_logged_in" not in st.session_state: st.session_state.is_logged_in = is_saved_logged_in

# ==============================================================================
# 2. FUNGSI UTAMA ENGINE AI & FILTER EKSTRAKSI DOKUMEN MURNI
# ==============================================================================
def muat_api_keys():
    """Membaca daftar API Keys dengan pemisahan baris baru (Anti-Gagal)."""
    keys = []
    
    # 1. Jalur Utama: Baca langsung per baris dari Environment Variable
    keys_env = os.getenv("GCP_API_KEYS")
    if keys_env:
        # Memisahkan text berdasarkan baris baru, lalu membersihkan spasi kosong
        for baris in keys_env.strip().split("\n"):
            baris_bersih = baris.strip().replace('"', '').replace(',', '').replace('[', '').replace(']', '')
            if baris_bersih and not baris_bersih.startswith("#"):
                keys.append(baris_bersih)
        if keys:
            return keys

    # 2. Jalur Cadangan (Fallback Lokal)
    if os.path.exists(NAMA_FILE_KEY):
        with open(NAMA_FILE_KEY, "r") as f:
            for baris in f:
                baris = baris.strip()
                if baris and not baris.startswith("#"):
                    keys.append(baris)
    return keys

def dapatkan_client_acak():
    daftar_keys = muat_api_keys()
    keys_valid = [k for k in daftar_keys if "Your" not in k and "Disini" not in k]
    if not keys_valid: return None
    return genai.Client(api_key=random.choice(keys_valid))

def ekstraks_dokumen_murni(teks_lengkap):
    """Fungsi pintar berlapis untuk memotong obrolan AI jika pembungkus bocor."""
    tag_mulai = "===Mulai Dokumen==="
    tag_akhir = "===Akhir Dokumen==="
    
    # Layer 1: Potong berdasarkan Tag Resmi jika ada
    if tag_mulai in teks_lengkap and tag_akhir in teks_lengkap:
        return teks_lengkap.split(tag_mulai)[1].split(tag_akhir)[0].strip()
    elif tag_mulai in teks_lengkap:
        return teks_lengkap.split(tag_mulai)[1].strip()
        
    # Layer 2 (Fallback): Jika AI lupa mencetak tag, bersihkan otomatis kalimat pengantar manual
    baris_teks = teks_lengkap.split('\n')
    baris_bersih = []
    mulai_rekam = False
    
    for b in baris_teks:
        if b.strip().startswith("#") or b.strip().startswith("**Bab") or b.strip().startswith("**Kuis") or b.strip().startswith("**Materi"):
            mulai_rekam = True
        if "Rekomendasi dan Tindak Lanjut" in b or "Sebagai seorang pakar" in b:
            continue
        if mulai_rekam:
            baris_bersih.append(b)
            
    if baris_bersih:
        return "\n".join(baris_bersih).strip()
        
    return teks_lengkap.strip()

def buat_file_docx(teks_markdown):
    doc = Document()
    for baris in teks_markdown.split('\n'):
        baris = baris.replace("===Mulai Dokumen===", "").replace("===Akhir Dokumen===", "")
        
        if baris.startswith('### '): doc.add_heading(baris.replace('### ', ''), level=3)
        elif baris.startswith('## '): doc.add_heading(baris.replace('## ', ''), level=2)
        elif baris.startswith('# '): doc.add_heading(baris.replace('# ', ''), level=1)
        elif baris.strip().startswith('- ') or baris.strip().startswith('* '):
            doc.add_paragraph(baris.strip()[2:], style='List Bullet')
        else:
            if baris.strip(): doc.add_paragraph(baris)
    buffer = io.BytesIO()
    doc.save(buffer)
    buffer.seek(0)
    return buffer

# ==============================================================================
# 3. INTERFACE LOGIN GOOGLE (ONE-CLICK SIGN IN SIMULATION)
# ==============================================================================
st.set_page_config(page_title="Universal AI Agent Pro", page_icon="🔮", layout="centered")

if not st.session_state.is_logged_in:
    st.write("")
    with st.container(border=True):
        st.markdown("<h2 style='text-align: center;'>🔮 Universal AI Agent Pro</h2>", unsafe_allow_html=True)
        st.markdown("<p style='text-align: center; color: gray;'>Masuk instan tanpa ribet ketik password</p>", unsafe_allow_html=True)
        st.markdown("---")
        st.write("🌐 **Pilih Akun Google Anda untuk Melanjutkan:**")
        logo_google_url = "https://upload.wikimedia.org/wikipedia/commons/c/c1/Google_%22G%22_logo.svg"
        col_g1, col_g2 = st.columns([1, 1])
        
        with col_g1:
            st.markdown(f"**<img src='{logo_google_url}' width='18'> Miftada Handoyo**", unsafe_allow_html=True)
            if st.button("handoyoyy1@gmail.com", use_container_width=True):
                st.session_state.user_email = EMAIL_ADMIN
                st.session_state.user_name = "Miftada Handoyo"
                st.session_state.is_logged_in = True
                simpan_session_login(EMAIL_ADMIN, "Miftada Handoyo")
                st.rerun()
                
        with col_g2:
            st.markdown(f"**<img src='{logo_google_url}' width='18'> User Pelanggan (Tester)**", unsafe_allow_html=True)
            if st.button("budi.santoso@gmail.com", use_container_width=True):
                st.session_state.user_email = "budi.santoso@gmail.com"
                st.session_state.user_name = "Budi Santoso"
                st.session_state.is_logged_in = True
                simpan_session_login("budi.santoso@gmail.com", "Budi Santoso")
                
                db_sekarang = muat_database_kv()
                if "budi.santoso@gmail.com" not in db_sekarang:
                    nominal_acak = 49000 + random.randint(1, 999)
                    db_sekarang["budi.santoso@gmail.com"] = {
                        "nama": "Budi Santoso", "status": "Pending Pembayaran",
                        "nominal_transfer": nominal_acak, "kode_aktivasi": f"PREM-{random.randint(100000, 999999)}"
                    }
                    simpan_database_kv(db_sekarang)
                st.rerun()
        st.markdown("---")
        
        with st.expander("🔐 Gunakan Akun Gmail Lainnya"):
            with st.form("custom_login"):
                custom_name = st.text_input("Nama Lengkap Anda:", placeholder="Contoh: Andi Wijaya")
                custom_email = st.text_input("Alamat Gmail Baru:", placeholder="contoh: andi.wijaya@gmail.com")
                submit_custom = st.form_submit_button("Masuk dengan Akun Baru 🚀", use_container_width=True)
                
                if submit_custom:
                    email_clean = custom_email.strip().lower()
                    if "@gmail.com" in email_clean and custom_name.strip() != "":
                        st.session_state.user_email = email_clean
                        st.session_state.user_name = custom_name.strip()
                        st.session_state.is_logged_in = True
                        simpan_session_login(email_clean, custom_name.strip())
                        
                        db_sekarang = muat_database_kv()
                        if email_clean not in db_sekarang:
                            nominal_acak = 49000 + random.randint(1, 999)
                            db_sekarang[email_clean] = {
                                "nama": custom_name.strip(), "status": "Pending Pembayaran",
                                "nominal_transfer": nominal_acak, "kode_aktivasi": f"PREM-{random.randint(100000, 999999)}"
                            }
                            simpan_database_kv(db_sekarang)
                        st.rerun()
                    else:
                        st.error("Gagal! Pastikan nama terisi dan domain email wajib menggunakan @gmail.com")
    st.stop()

email_aktif = st.session_state.user_email
nama_aktif = st.session_state.user_name
status_user = st.session_state.db_users[email_aktif]["status"] if email_aktif in st.session_state.db_users else "Pending Pembayaran"

col_profil, col_nav_out = st.columns([4, 1])
with col_profil:
    st.write(f"👋 Halo, **{nama_aktif}** (`{email_aktif}`) | Lisensi: **{status_user}**")
with col_nav_out:
    if st.button("Keluar 🚪", use_container_width=True):
        hapus_session_login()
        st.session_state.is_logged_in = False
        st.session_state.user_email = ""
        st.session_state.user_name = ""
        st.rerun()

st.markdown("---")

# ==============================================================================
# 4. HALAMAN LOCK SCREEN PEMBAYARAN
# ==============================================================================
if status_user == "Pending Pembayaran" and email_aktif != EMAIL_ADMIN:
    st.title("🔒 Akses Layanan Premium Terkunci")
    nominal_bayar = st.session_state.db_users[email_aktif]["nominal_transfer"]
    with st.container(border=True):
        st.markdown("### 💳 Prosedur Aktivasi Premium")
        st.subheader(f"Rp {nominal_bayar:,}")
        col_rek, col_gambar_qris = st.columns([1, 1])
        with col_rek:
            st.markdown("**Tujuan Pembayaran:**\n* **E-Wallet (Gopay):** `081333260292`\n* **Atas Nama:** Miftada Handoyo Y.")
            if st.button("🔄 Cek Status Aktivasi Saya", use_container_width=True):
                st.session_state.db_users = muat_database_kv()
                st.rerun()
        with col_gambar_qris:
            if os.path.exists("qris.jpeg"): st.image("qris.jpeg", use_container_width=True)
    st.stop()

# ==============================================================================
# 5. HALAMAN UTAMA WORKSPACE & KONTROL TAB
# ==============================================================================
if email_aktif == EMAIL_ADMIN:
    tab_workspace, tab_settings, tab_admin = st.tabs(["🚀 Ruang Kerja Agent", "⚙️ Konfigurasi API", "👑 Dashboard Admin"])
else:
    tab_workspace, tab_settings = st.tabs(["🚀 Ruang Kerja Agent", "⚙️ Informasi Sistem"])

# --- TAB 1: WORKSPACE ---
with tab_workspace:
    st.write("")
    with st.container(border=True):
        st.markdown("### ✨ Spesifikasi Kebutuhan Kerja")
        profesi_user = st.text_input("Tentukan Peran Ahli / Profesi Spesialis Terfokus:", placeholder="Contoh: Senior Database Administrator, Kepala Sekolah, Guru Matematika SD")
        tugas_user = st.text_area("Deskripsikan Masalah Teknis atau Tugas yang Harus Diselesaikan:", placeholder="Salin perintah pengerjaan berkas di sini...", height=150)
        
    st.write("")
    tombol_proses = st.button("Jalankan Mesin Otomasi Agen ⚡", type="primary", use_container_width=True)

    if "hasil_ai" not in st.session_state: st.session_state.hasil_ai = ""

    if tombol_proses:
        if profesi_user.strip() != "" and tugas_user.strip() != "":
            client = dapatkan_client_acak()
            if client is not None:
                with st.spinner(f"Mesin Otomasi sedang memproses solusi sebagai {profesi_user}..."):
                    try:
                        instruksi_kepribadian = (
                            f"Anda adalah seorang pakar top dunia dan profesional senior di bidang: '{profesi_user}'. "
                            f"Jawablah tugas dari user dengan kualitas industri tertinggi. "
                            f"PENTING & WAJIB: Anda harus memisahkan antara bagian teks obrolan sapaan dengan isi dokumen utama. "
                            f"Bungkus dokumen inti buatan Anda dengan menggunakan penanda teks eksak berikut:\n"
                            f"===Mulai Dokumen===\n[Isi dokumen/materi utama Anda di sini]\n===Akhir Dokumen===\n"
                            f"Jangan biarkan ada teks pengantar atau penutup bawaan AI masuk ke dalam tanda pembungkus tersebut."
                        )
                        
                        respons_final = client.models.generate_content(
                            model='gemini-2.5-flash',
                            contents=tugas_user,
                            config=types.GenerateContentConfig(
                                system_instruction=instruksi_kepribadian,
                                temperature=0.0  # Memaksa model patuh 100% pada struktur penanda tag
                            ),
                        )
                        st.session_state.hasil_ai = respons_final.text
                        st.balloons()
                    except Exception as e:
                        st.error(f"Terjadi interupsi jaringan API: {e}")
        else:
            st.warning("Kolom Isian Peran Ahli dan Detail Masalah wajib diisi.")

    if st.session_state.hasil_ai != "":
        st.write("")
        st.markdown(f"### 📋 Dokumentasi Pratinjau Sistem: *{profesi_user}*")
        
        with st.container(border=True): 
            st.markdown(st.session_state.hasil_ai)
            
        # Proses Filter Otomatis
        dokumen_murni_unduhan = ekstraks_dokumen_murni(st.session_state.hasil_ai)
        
        st.write("")
        with st.expander("💾 Opsi Penyimpanan Berkas Bersih (Rapi & Tanpa Penjelasan AI)", expanded=True):
            col_word, col_txt = st.columns(2)
            with col_word:
                st.download_button(
                    label="📝 Unduh Dokumen Word (.docx) Bersih", 
                    data=buat_file_docx(dokumen_murni_unduhan), 
                    file_name=f"Clean_{profesi_user.replace(' ', '_')}.docx", 
                    mime="application/vnd.openxmlformats-officedocument.wordprocessingml.document", 
                    use_container_width=True
                )
            with col_txt:
                st.download_button(
                    label="📄 Unduh Teks Mentah (.txt) Bersih", 
                    data=dokumen_murni_unduhan, 
                    file_name=f"Clean_{profesi_user.replace(' ', '_')}.txt", 
                    mime="text/plain", 
                    use_container_width=True
                )

# --- TAB 2: SYSTEM CONFIG ---
with tab_settings:
    st.write("")
    st.markdown("### ⚙️ Keseimbangan Token Pool")
    keys_aktif = [k for k in muat_api_keys() if "Your" not in k and "Disini" not in k]
    st.metric(label="Jumlah Kunci Cadangan API Aktif", value=len(keys_aktif))

# --- TAB 3: CONTROL PANEL ADMIN (TERISOLASI) ---
if email_aktif == EMAIL_ADMIN:
    with tab_admin:
        st.write("")
        st.markdown("### 👑 Panel Manajemen Aktivasi Manual")
        if st.button("🔄 Refresh Data Pengajuan Masuk", use_container_width=True):
            st.session_state.db_users = muat_database_kv()
            st.rerun()
            
        db_sekarang = muat_database_kv()
        ada_user = False
        for user_mail, data in list(db_sekarang.items()):
            if user_mail == EMAIL_ADMIN: continue
            ada_user = True
            with st.container(border=True):
                col_data_user, col_tombol_konfirmasi = st.columns([3, 1])
                with col_data_user:
                    st.markdown(f"👤 **Nama:** {data.get('nama', 'User Baru')} | `{user_mail}`")
                    st.write(f"💰 Tagihan: **Rp {data['nominal_transfer']:,}** | Status: `{data['status']}`")
                with col_tombol_konfirmasi:
                    if data['status'] == "Pending Pembayaran":
                        if st.button("Confirm 🟩", key=f"adm_act_{user_mail}", use_container_width=True):
                            db_sekarang = muat_database_kv()
                            db_sekarang[user_mail]["status"] = "Aktif"
                            simpan_database_kv(db_sekarang)
                            st.session_state.db_users = db_sekarang
                            st.rerun()
                    else:
                        st.button("Aktif ✅", key=f"adm_done_{user_mail}", disabled=True, use_container_width=True)
                        
        if not ada_user:
            st.info("Belum ada data pelanggan terdaftar.")
            
        # ==============================================================================
        # INSPEKTUR DATA MURNI LOKAL / KV STORAGE (OTOMATIS MENYESUAIKAN)
        # ==============================================================================
        st.write("")
        st.markdown("---")
        st.markdown("### 🔍 Inspektur Data Mentah (Mode Dev/Admin)")
        if st.button("Lihat Seluruh Isi DB Murni 📂", use_container_width=True):
            data_inspeksi = muat_database_kv()
            st.json(data_inspeksi)

st.write("")
st.markdown("---")
st.caption("© 2026 Universal AI Agent - Commercial Framework Pro Edition")
